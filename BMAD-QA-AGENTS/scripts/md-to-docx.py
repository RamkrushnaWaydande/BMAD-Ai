#!/usr/bin/env python3
"""
Markdown to DOCX converter for SRS and other agent outputs.

Usage:
    python3 scripts/md-to-docx.py <input.md> [output.docx]

If output path is not provided, the .docx file is saved alongside the input
with the same name but a .docx extension.

Requires: python-docx (pip install python-docx)
"""

import sys
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def parse_metadata(lines):
    """Extract metadata key-value pairs from the top of the markdown."""
    metadata = {}
    for line in lines:
        line = line.strip()
        if line.startswith("**") and ":**" in line:
            match = re.match(r"\*\*(.+?):\*\*\s*(.+)", line)
            if match:
                metadata[match.group(1).strip()] = match.group(2).strip()
        elif line == "---":
            break
    return metadata


def add_styled_paragraph(doc, text, style_name, bold=False, alignment=None):
    """Add a paragraph with the given style."""
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    run.bold = bold
    if alignment:
        para.alignment = alignment
    return para


def parse_table(lines, start_idx):
    """Parse a markdown table starting at the given index. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        # Skip separator rows (|---|---|)
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def add_table_to_doc(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_text.strip()
                # Bold the header row
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def process_inline_formatting(paragraph, text):
    """Process bold, italic, and code formatting within a line."""
    # Split text by formatting markers and add runs accordingly
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def convert_md_to_docx(input_path, output_path=None):
    """Convert a markdown file to a DOCX document."""
    input_path = Path(input_path)
    if output_path is None:
        output_path = input_path.with_suffix(".docx")
    else:
        output_path = Path(output_path)

    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Configure heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = RGBColor(0, 0, 0)

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_lines)
                para = doc.add_paragraph(style="Normal")
                run = para.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip horizontal rules
        if stripped == "---":
            i += 1
            continue

        # Headings
        if stripped.startswith("#"):
            match = re.match(r"^(#{1,6})\s+(.*)", stripped)
            if match:
                level = min(len(match.group(1)), 4)  # DOCX supports up to Heading 9, but keep it reasonable
                heading_text = match.group(2).strip()
                # Remove any markdown formatting from heading
                heading_text = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_text)
                heading_text = re.sub(r"\*(.+?)\*", r"\1", heading_text)
                doc.add_heading(heading_text, level=level)
                i += 1
                continue

        # Metadata lines (bold key: value)
        if stripped.startswith("**") and ":**" in stripped:
            para = doc.add_paragraph()
            match = re.match(r"\*\*(.+?):\*\*\s*(.*)", stripped)
            if match:
                run = para.add_run(f"{match.group(1)}: ")
                run.bold = True
                para.add_run(match.group(2))
            else:
                process_inline_formatting(para, stripped)
            i += 1
            continue

        # Tables
        if stripped.startswith("|"):
            rows, end_idx = parse_table(lines, i)
            add_table_to_doc(doc, rows)
            i = end_idx
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            # Determine indentation level
            indent_level = (len(line) - len(line.lstrip())) // 2
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            process_inline_formatting(para, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r"^(\s*)(\d+)\.\s+(.*)", line)
        if num_match:
            indent_level = len(num_match.group(1)) // 2
            text = num_match.group(3)
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            process_inline_formatting(para, text)
            i += 1
            continue

        # Blockquotes
        if stripped.startswith("> "):
            text = stripped[2:]
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            process_inline_formatting(para, text)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraphs
        para = doc.add_paragraph()
        process_inline_formatting(para, stripped)
        i += 1

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"DOCX saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/md-to-docx.py <input.md> [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.exists(input_file):
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)

    convert_md_to_docx(input_file, output_file)
